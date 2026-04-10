"""WordExporter + HTMLRenderer + 模板加载 单元测试。"""
import asyncio
import io
import tempfile

import pytest
import pytest_asyncio
from docx import Document

from sololab.modules.writer.export.word_exporter import WordExporter
from sololab.modules.writer.export.html_renderer import render_document_html
from sololab.modules.writer.templates.registry import TemplateRegistry
from pathlib import Path


@pytest.fixture
def sample_doc():
    """完整的示例文档。"""
    return {
        "doc_id": "test-doc-export",
        "session_id": "test-session",
        "title": "Vision Transformers for Medical Image Segmentation",
        "template_id": "nature",
        "language": "en",
        "status": "complete",
        "sections": [
            {
                "id": "s1", "type": "abstract", "title": "Abstract",
                "content": "<p>This paper presents a novel approach to medical image segmentation using Vision Transformers.</p>",
                "order": 0, "status": "complete", "word_count": 15,
            },
            {
                "id": "s2", "type": "introduction", "title": "Introduction",
                "content": "<p>Medical image segmentation is a critical task in clinical practice [1]. Recent advances in Vision Transformers [2] have shown promising results.</p><p>In this work, we propose a new architecture that combines CNNs with ViT.</p>",
                "order": 1, "status": "complete", "word_count": 35,
            },
            {
                "id": "s3", "type": "methods", "title": "Methods",
                "content": "<p>Our method consists of three components:</p><ul><li>A CNN encoder for local features</li><li>A ViT module for global attention</li><li>A fusion decoder</li></ul>",
                "order": 2, "status": "complete", "word_count": 25,
            },
        ],
        "references": [
            {"number": 1, "title": "U-Net: Convolutional Networks for Biomedical Image Segmentation", "authors": ["Ronneberger, O.", "Fischer, P.", "Brox, T."], "year": 2015, "venue": "MICCAI"},
            {"number": 2, "title": "An Image is Worth 16x16 Words", "authors": ["Dosovitskiy, A.", "Beyer, L.", "Kolesnikov, A."], "year": 2021, "venue": "ICLR"},
        ],
        "figures": [
            {"id": "fig1", "section_id": "s3", "caption": "Architecture overview", "url": "/storage/writer/test/fig1.png", "order": 1, "number": 1},
        ],
        "metadata": {},
        "word_count": 75,
    }


@pytest.fixture
def minimal_doc():
    return {
        "doc_id": "min",
        "title": "Minimal Paper",
        "template_id": "nature",
        "sections": [{"id": "s1", "type": "abstract", "title": "Abstract", "content": "<p>Hello world.</p>", "order": 0, "status": "complete", "word_count": 2}],
        "references": [],
        "figures": [],
    }


# ── WordExporter Tests ──────────────────────────────────

class TestWordExporter:
    """测试 Word 导出。"""

    def test_export_produces_valid_docx(self, sample_doc):
        exporter = WordExporter(storage_path=tempfile.mkdtemp())
        docx_bytes = asyncio.run(
            exporter.export(sample_doc, "nature-numeric")
        )
        assert len(docx_bytes) > 0
        # Verify it's a valid docx
        doc = Document(io.BytesIO(docx_bytes))
        assert len(doc.paragraphs) > 0

    def test_export_contains_title(self, sample_doc):
        exporter = WordExporter(storage_path=tempfile.mkdtemp())
        docx_bytes = asyncio.run(exporter.export(sample_doc))
        doc = Document(io.BytesIO(docx_bytes))
        texts = [p.text for p in doc.paragraphs]
        assert any("Vision Transformers" in t for t in texts)

    def test_export_contains_sections(self, sample_doc):
        exporter = WordExporter(storage_path=tempfile.mkdtemp())
        docx_bytes = asyncio.run(exporter.export(sample_doc))
        doc = Document(io.BytesIO(docx_bytes))
        texts = " ".join(p.text for p in doc.paragraphs)
        assert "Abstract" in texts
        assert "Introduction" in texts
        assert "Methods" in texts

    def test_export_contains_references(self, sample_doc):
        exporter = WordExporter(storage_path=tempfile.mkdtemp())
        docx_bytes = asyncio.run(exporter.export(sample_doc))
        doc = Document(io.BytesIO(docx_bytes))
        texts = " ".join(p.text for p in doc.paragraphs)
        assert "References" in texts
        assert "U-Net" in texts

    def test_export_html_list_items(self, sample_doc):
        """HTML <ul><li> 应转为 Word 段落。"""
        exporter = WordExporter(storage_path=tempfile.mkdtemp())
        docx_bytes = asyncio.run(exporter.export(sample_doc))
        doc = Document(io.BytesIO(docx_bytes))
        texts = " ".join(p.text for p in doc.paragraphs)
        assert "CNN encoder" in texts
        assert "ViT module" in texts

    def test_export_minimal_doc(self, minimal_doc):
        exporter = WordExporter(storage_path=tempfile.mkdtemp())
        docx_bytes = asyncio.run(exporter.export(minimal_doc))
        assert len(docx_bytes) > 0
        doc = Document(io.BytesIO(docx_bytes))
        assert len(doc.paragraphs) > 0

    def test_export_with_ieee_style(self, sample_doc):
        sample_doc["template_id"] = "ieee"
        exporter = WordExporter(storage_path=tempfile.mkdtemp())
        docx_bytes = asyncio.run(
            exporter.export(sample_doc, "ieee-numeric")
        )
        doc = Document(io.BytesIO(docx_bytes))
        texts = " ".join(p.text for p in doc.paragraphs)
        assert "U-Net" in texts  # Reference still present

    def test_export_empty_sections(self):
        doc = {"doc_id": "e", "title": "Empty", "template_id": "nature", "sections": [], "references": [], "figures": []}
        exporter = WordExporter(storage_path=tempfile.mkdtemp())
        docx_bytes = asyncio.run(exporter.export(doc))
        assert len(docx_bytes) > 0


# ── HTML Renderer Tests ─────────────────────────────────

class TestHTMLRenderer:
    """测试 HTML 预览渲染。"""

    def test_renders_title(self, sample_doc):
        html = render_document_html(sample_doc)
        assert "Vision Transformers" in html

    def test_renders_sections(self, sample_doc):
        html = render_document_html(sample_doc)
        assert "Abstract" in html
        assert "Introduction" in html
        assert "Methods" in html

    def test_renders_references(self, sample_doc):
        html = render_document_html(sample_doc, "nature-numeric")
        assert "U-Net" in html
        assert "reference-list" in html

    def test_renders_figures(self, sample_doc):
        html = render_document_html(sample_doc)
        assert "paper-figure" in html
        assert "Architecture overview" in html

    def test_renders_cvpr_style(self, sample_doc):
        sample_doc["template_id"] = "cvpr"
        html = render_document_html(sample_doc)
        assert "column-count: 2" in html

    def test_renders_chinese_style(self, sample_doc):
        sample_doc["template_id"] = "chinese_journal"
        html = render_document_html(sample_doc)
        assert "SimSun" in html or "text-indent" in html

    def test_renders_empty_section_placeholder(self):
        doc = {
            "title": "T", "template_id": "nature",
            "sections": [{"id": "s1", "type": "abstract", "title": "Abstract", "content": "", "status": "empty"}],
            "references": [], "figures": [],
        }
        html = render_document_html(doc)
        assert "placeholder" in html


# ── Template Loading Tests ──────────────────────────────

class TestAllTemplatesLoad:
    """验证所有 YAML 模板都能正确加载。"""

    def test_load_all_templates(self):
        templates_dir = Path(__file__).resolve().parents[5] / "backend" / "src" / "sololab" / "modules" / "writer" / "templates"
        registry = TemplateRegistry(templates_dir)
        templates = registry.list_all()
        ids = registry.list_ids()

        assert len(templates) >= 6  # nature + cvpr + iccv + ieee + acm + chinese_journal
        assert "nature" in ids
        assert "cvpr" in ids
        assert "iccv" in ids
        assert "ieee" in ids
        assert "acm" in ids
        assert "chinese_journal" in ids

    def test_cvpr_template_details(self):
        templates_dir = Path(__file__).resolve().parents[5] / "backend" / "src" / "sololab" / "modules" / "writer" / "templates"
        registry = TemplateRegistry(templates_dir)
        t = registry.get("cvpr")
        assert t is not None
        assert t.page_limit == 8
        assert t.citation.style == "ieee-numeric"
        section_types = [s.type for s in t.sections]
        assert "related_work" in section_types
        assert "experiments" in section_types

    def test_chinese_journal_template(self):
        templates_dir = Path(__file__).resolve().parents[5] / "backend" / "src" / "sololab" / "modules" / "writer" / "templates"
        registry = TemplateRegistry(templates_dir)
        t = registry.get("chinese_journal")
        assert t is not None
        assert t.language_default == "zh"
        assert t.citation.style == "gbt-7714"
        assert t.citation.max_authors == 3
        section_types = [s.type for s in t.sections]
        assert "abstract_zh" in section_types
        assert "abstract_en" in section_types

    def test_acm_template(self):
        templates_dir = Path(__file__).resolve().parents[5] / "backend" / "src" / "sololab" / "modules" / "writer" / "templates"
        registry = TemplateRegistry(templates_dir)
        t = registry.get("acm")
        assert t is not None
        assert t.citation.style == "apa-author-year"
        assert t.page_limit == 12
