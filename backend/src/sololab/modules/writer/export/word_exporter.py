"""WordExporter — 将 WriterAI 文档导出为 .docx Word 文件。

基于 python-docx，支持：
- 模板样式加载（.docx 模板文件）
- HTML 章节内容 → Word 段落映射（含 <table>、<h1-3>）
- LaTeX 公式 → Unicode 简易渲染
- 中文字体配置（East Asian font）
- 图表嵌入（URL 下载 + Caption + 编号）
- 参考文献列表（按模板引用格式）
"""
from __future__ import annotations

import io
import logging
import os
import re
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from sololab.modules.writer.export.citation_formatter import format_reference

logger = logging.getLogger(__name__)

TEMPLATES_DIR = Path(__file__).parent.parent / "templates" / "docx"

# ── LaTeX → Unicode 简易转换表 ────────────────────────
_LATEX_UNICODE = {
    # 小写希腊字母
    r"\alpha": "α", r"\beta": "β", r"\gamma": "γ", r"\delta": "δ",
    r"\epsilon": "ε", r"\varepsilon": "ε", r"\zeta": "ζ", r"\eta": "η",
    r"\theta": "θ", r"\vartheta": "ϑ", r"\iota": "ι", r"\kappa": "κ",
    r"\lambda": "λ", r"\mu": "μ", r"\nu": "ν", r"\xi": "ξ",
    r"\pi": "π", r"\varpi": "ϖ", r"\rho": "ρ", r"\varrho": "ϱ",
    r"\sigma": "σ", r"\varsigma": "ς", r"\tau": "τ", r"\upsilon": "υ",
    r"\phi": "φ", r"\varphi": "φ", r"\chi": "χ", r"\psi": "ψ", r"\omega": "ω",
    # 大写希腊字母
    r"\Gamma": "Γ", r"\Delta": "Δ", r"\Theta": "Θ", r"\Lambda": "Λ",
    r"\Xi": "Ξ", r"\Pi": "Π", r"\Sigma": "Σ", r"\Upsilon": "Υ",
    r"\Phi": "Φ", r"\Psi": "Ψ", r"\Omega": "Ω",
    # 运算符与符号
    r"\sum": "∑", r"\prod": "∏", r"\int": "∫", r"\oint": "∮",
    r"\partial": "∂", r"\nabla": "∇", r"\infty": "∞",
    r"\pm": "±", r"\mp": "∓", r"\times": "×", r"\div": "÷",
    r"\cdot": "·", r"\cdots": "⋯", r"\ldots": "…", r"\dots": "…",
    r"\leq": "≤", r"\le": "≤", r"\geq": "≥", r"\ge": "≥",
    r"\neq": "≠", r"\ne": "≠", r"\approx": "≈", r"\equiv": "≡",
    r"\sim": "∼", r"\propto": "∝", r"\ll": "≪", r"\gg": "≫",
    r"\subset": "⊂", r"\supset": "⊃", r"\subseteq": "⊆", r"\supseteq": "⊇",
    r"\in": "∈", r"\notin": "∉", r"\ni": "∋",
    r"\cup": "∪", r"\cap": "∩", r"\emptyset": "∅", r"\varnothing": "∅",
    r"\forall": "∀", r"\exists": "∃", r"\nexists": "∄",
    r"\rightarrow": "→", r"\leftarrow": "←", r"\leftrightarrow": "↔",
    r"\Rightarrow": "⇒", r"\Leftarrow": "⇐", r"\Leftrightarrow": "⇔",
    r"\to": "→", r"\mapsto": "↦", r"\implies": "⟹", r"\iff": "⟺",
    r"\langle": "⟨", r"\rangle": "⟩", r"\hbar": "ℏ", r"\ell": "ℓ",
    r"\Re": "ℜ", r"\Im": "ℑ", r"\aleph": "ℵ",
    # 常见函数（去掉反斜杠）
    r"\sin": "sin", r"\cos": "cos", r"\tan": "tan", r"\cot": "cot",
    r"\sec": "sec", r"\csc": "csc",
    r"\log": "log", r"\ln": "ln", r"\exp": "exp",
    r"\max": "max", r"\min": "min", r"\sup": "sup", r"\inf": "inf",
    r"\arg": "arg", r"\det": "det", r"\dim": "dim", r"\ker": "ker",
}

# 上标 / 下标的常见 Unicode 映射（仅数字和少量字符）
_SUPERSCRIPT = str.maketrans("0123456789+-=()n", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁿ")
_SUBSCRIPT = str.maketrans("0123456789+-=()", "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎")


def _latex_to_unicode(latex: str) -> str:
    """Convert simple LaTeX to readable Unicode text. Best-effort, not full parser."""
    s = latex.strip()
    # 去掉常见包装：\text{...} \mathbf{...} 等
    for cmd in ("text", "mathbf", "mathit", "mathrm", "mathbb", "mathcal", "mathsf", "mathtt", "operatorname"):
        s = re.sub(r"\\" + cmd + r"\s*\{([^{}]*)\}", r"\1", s)
    # \frac{a}{b} → (a)/(b)
    s = re.sub(r"\\frac\s*\{([^{}]+)\}\s*\{([^{}]+)\}", r"(\1)/(\2)", s)
    # \sqrt[n]{a} → ⁿ√(a)；\sqrt{a} → √(a)
    s = re.sub(r"\\sqrt\s*\[([^\]]+)\]\s*\{([^{}]+)\}", r"\1√(\2)", s)
    s = re.sub(r"\\sqrt\s*\{([^{}]+)\}", r"√(\1)", s)
    # 替换符号表（按长度降序，避免短前缀误匹配）
    for cmd, ch in sorted(_LATEX_UNICODE.items(), key=lambda x: -len(x[0])):
        s = s.replace(cmd, ch)
    # 上下标：x_2 → x₂，x^2 → x²；{...} 形式简单处理单字符
    s = re.sub(r"\^\{([^{}]+)\}", lambda m: m.group(1).translate(_SUPERSCRIPT), s)
    s = re.sub(r"_\{([^{}]+)\}", lambda m: m.group(1).translate(_SUBSCRIPT), s)
    s = re.sub(r"\^(.)", lambda m: m.group(1).translate(_SUPERSCRIPT), s)
    s = re.sub(r"_(.)", lambda m: m.group(1).translate(_SUBSCRIPT), s)
    # 去掉剩余未识别命令的反斜杠
    s = re.sub(r"\\([a-zA-Z]+)", r"\1", s)
    # 清理大括号
    s = s.replace("{", "").replace("}", "")
    return s


def _split_latex_segments(text: str) -> list[tuple[str, bool]]:
    """Split text into (segment, is_math) tuples by $...$, $$...$$, \\(...\\), \\[...\\]."""
    pattern = re.compile(
        r"(\$\$[\s\S]+?\$\$|\\\[[\s\S]+?\\\]|\$[^\$\n]+?\$|\\\([^)]+?\\\))"
    )
    segments: list[tuple[str, bool]] = []
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            segments.append((text[pos:m.start()], False))
        raw = m.group(0)
        if raw.startswith("$$"):
            inner = raw[2:-2]
        elif raw.startswith("\\["):
            inner = raw[2:-2]
        elif raw.startswith("\\("):
            inner = raw[2:-2]
        else:
            inner = raw[1:-1]
        segments.append((_latex_to_unicode(inner), True))
        pos = m.end()
    if pos < len(text):
        segments.append((text[pos:], False))
    return segments or [(text, False)]


# ── 字体配置帮助函数 ────────────────────────────────
def _set_eastasia_font(run_or_style, font_name: str = "SimSun") -> None:
    """Set the East Asian font for a Run or Style (CJK character rendering)."""
    element = run_or_style._element if hasattr(run_or_style, "_element") else run_or_style.element
    rPr = element.get_or_add_rPr() if hasattr(element, "get_or_add_rPr") else element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:eastAsiaTheme"), "")  # clear theme override


def _setup_default_styles(document: Document) -> None:
    """Set up basic styles for a blank document with proper CJK font support."""
    # Normal style
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.4
    _set_eastasia_font(style, "SimSun")  # 宋体为通用 CJK fallback

    # Heading styles — override default blue with black
    for level in range(1, 4):
        try:
            h_style = document.styles[f"Heading {level}"]
            h_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            h_style.font.name = "Times New Roman"
            h_style.font.bold = True
            sizes = {1: 16, 2: 13, 3: 11}
            h_style.font.size = Pt(sizes.get(level, 11))
            _set_eastasia_font(h_style, "SimHei")  # 标题用黑体
        except KeyError:
            pass


# ── 主导出器 ─────────────────────────────────────────
class WordExporter:
    """Export WriterAI documents to .docx files."""

    def __init__(self, storage_path: str = "./storage") -> None:
        self.storage_path = Path(storage_path)

    async def export(self, doc: dict, citation_style: str = "nature-numeric") -> bytes:
        template_id = doc.get("template_id", "nature")
        template_path = TEMPLATES_DIR / f"{template_id}.docx"

        if template_path.exists():
            document = Document(str(template_path))
        else:
            document = Document()
        _setup_default_styles(document)

        # Title
        title = doc.get("title", "Untitled Paper")
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        _set_eastasia_font(run, "SimHei")

        document.add_paragraph()

        # Sections
        for section in doc.get("sections", []):
            self._add_section(document, section, doc)

        # Global figures (no section_id)
        for fig in doc.get("figures", []):
            if not fig.get("section_id"):
                self._add_figure(document, fig)

        # References
        refs = doc.get("references", [])
        if refs:
            document.add_paragraph()
            self._add_heading(document, "参考文献" if _has_cjk(title) else "References", level=1)
            for ref in refs:
                formatted = format_reference(ref, citation_style)
                p = document.add_paragraph(f"[{ref.get('number', '')}] {formatted}")
                p.paragraph_format.first_line_indent = Pt(-18)
                p.paragraph_format.left_indent = Pt(18)
                for run in p.runs:
                    run.font.size = Pt(9)
                    _set_eastasia_font(run, "SimSun")

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _add_heading(self, document: Document, text: str, level: int = 1) -> None:
        """Add a heading with proper CJK font and black color."""
        h = document.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            _set_eastasia_font(run, "SimHei")

    def _add_section(self, document: Document, section: dict, doc: dict) -> None:
        title = section.get("title", "")
        content = section.get("content", "")
        section_id = section.get("id", "")

        if title:
            self._add_heading(document, title, level=1)

        if content:
            parser = _HTMLToDocxParser(document)
            parser.feed(content)
            parser.close()

        for fig in doc.get("figures", []):
            if fig.get("section_id") == section_id:
                self._add_figure(document, fig)

    def _add_figure(self, document: Document, fig: dict) -> None:
        url = fig.get("url", "")
        caption = fig.get("caption", "")
        number = fig.get("order", fig.get("number", 0))

        img_path = self._resolve_figure_path(url)
        if img_path and os.path.exists(img_path):
            try:
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(5.0))
            except Exception:
                logger.warning("Failed to embed figure: %s", img_path)
                document.add_paragraph(f"[Figure {number}: image not available]")
        else:
            document.add_paragraph(f"[Figure {number}: {url}]")

        if caption:
            cap_p = document.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap_p.add_run(f"图 {number}. {caption}" if _has_cjk(caption) else f"Figure {number}. {caption}")
            run.italic = True
            run.font.size = Pt(9)
            _set_eastasia_font(run, "SimSun")

    def _resolve_figure_path(self, url: str) -> str | None:
        if not url:
            return None
        if url.startswith("/storage/"):
            return str(self.storage_path / url[len("/storage/"):])
        if os.path.isabs(url):
            return url
        return None


def _has_cjk(text: str) -> bool:
    """Check if string contains any CJK characters."""
    return any("\u4e00" <= c <= "\u9fff" for c in text)


# ── HTML → docx 解析器 ────────────────────────────────
class _HTMLToDocxParser(HTMLParser):
    """HTML → python-docx 转换器。

    支持：<p> <strong>/<b> <em>/<i> <ul>/<ol>/<li> <br>
         <h1>-<h3> <table>/<thead>/<tbody>/<tr>/<td>/<th>
         LaTeX 公式 ($...$ / $$...$$ / \\(\\) / \\[\\])
    """

    def __init__(self, document: Document) -> None:
        super().__init__()
        self.document = document
        self._current_paragraph = None
        self._bold = False
        self._italic = False
        self._heading_level: int | None = None
        # List state
        self._in_list = False
        self._list_ordered = False
        self._list_counter = 0
        # Table state
        self._table = None
        self._current_row_cells: list = []
        self._table_col_count = 0
        self._current_cell = None
        self._in_thead = False

    def handle_starttag(self, tag: str, attrs: list) -> None:
        tag = tag.lower()
        if tag in ("p", "div"):
            self._current_paragraph = self.document.add_paragraph()
        elif tag in ("strong", "b"):
            self._bold = True
        elif tag in ("em", "i"):
            self._italic = True
        elif tag == "br":
            if self._current_paragraph:
                self._current_paragraph.add_run("\n")
        elif tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._heading_level = int(tag[1])
            self._current_paragraph = self.document.add_heading("", level=min(self._heading_level, 3))
        elif tag == "ul":
            self._in_list = True
            self._list_ordered = False
        elif tag == "ol":
            self._in_list = True
            self._list_ordered = True
            self._list_counter = 0
        elif tag == "li":
            self._list_counter += 1
            prefix = f"{self._list_counter}. " if self._list_ordered else "• "
            self._current_paragraph = self.document.add_paragraph(prefix)
        elif tag == "table":
            self._table = self.document.add_table(rows=0, cols=0)
            try:
                self._table.style = "Light Grid Accent 1"
            except KeyError:
                pass
            self._table_col_count = 0
        elif tag == "thead":
            self._in_thead = True
        elif tag == "tbody":
            self._in_thead = False
        elif tag == "tr":
            if self._table is not None:
                # 第一行决定列数
                if self._table_col_count == 0:
                    # 推迟到 td/th 时再加 row
                    self._current_row_cells = []
                else:
                    row = self._table.add_row()
                    self._current_row_cells = list(row.cells)
        elif tag in ("td", "th"):
            if self._table is None:
                return
            if self._table_col_count == 0:
                # 第一行：动态扩列
                self._table.add_column(Inches(1.0))
                self._table_col_count = len(self._table.columns)
                if not self._table.rows:
                    row = self._table.add_row()
                    self._current_row_cells = list(row.cells)
                else:
                    self._current_row_cells = list(self._table.rows[0].cells)
                cell_idx = self._table_col_count - 1
            else:
                # 后续行：当前 row 的下一个空 cell
                cell_idx = sum(1 for c in self._current_row_cells if c.text)
                if cell_idx >= len(self._current_row_cells):
                    return
            self._current_cell = self._current_row_cells[cell_idx] if cell_idx < len(self._current_row_cells) else None
            self._current_paragraph = self._current_cell.paragraphs[0] if self._current_cell else None
            if tag == "th":
                self._bold = True

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in ("strong", "b"):
            self._bold = False
        elif tag in ("em", "i"):
            self._italic = False
        elif tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            # Apply heading styling
            if self._current_paragraph is not None:
                for run in self._current_paragraph.runs:
                    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
                    _set_eastasia_font(run, "SimHei")
            self._heading_level = None
            self._current_paragraph = None
        elif tag in ("ul", "ol"):
            self._in_list = False
            self._list_counter = 0
        elif tag in ("p", "div", "li"):
            self._current_paragraph = None
        elif tag == "table":
            self._table = None
            self._current_row_cells = []
            self._table_col_count = 0
            self._current_cell = None
        elif tag in ("td", "th"):
            self._current_cell = None
            self._current_paragraph = None
            if tag == "th":
                self._bold = False

    def handle_data(self, data: str) -> None:
        text = data.strip("\n")
        if not text.strip():
            return
        if self._current_paragraph is None:
            self._current_paragraph = self.document.add_paragraph()

        # Process LaTeX segments
        for segment, is_math in _split_latex_segments(text):
            if not segment:
                continue
            run = self._current_paragraph.add_run(segment)
            run.bold = self._bold
            run.italic = self._italic or is_math
            run.font.size = Pt(11)
            _set_eastasia_font(run, "SimSun")
